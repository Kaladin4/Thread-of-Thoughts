from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
import os
from langchain_core.documents.base import Document
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from langchain.schema import Document
def split_data(data):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    splits = text_splitter.split_documents(data)
    return splits
def get_retriever(memory_to_load):
    vectorstore = FAISS.from_documents(
        documents=memory_to_load,
        embedding=OpenAIEmbeddings(
            model="embeddings",
        ),
    )
    retriever = vectorstore.as_retriever(search_type="similarity")
    return retriever
def format_docs(docs):
        return "\n\n".join(doc.page_content for doc in docs)




def read_file(file_path):
    """
    Reads the content of various file types and returns it as a string.
    
    Args:
        file_path (str): The path to the file.
    
    Returns:
        str: The content of the file.
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        elif file_extension == '.pdf':
            text = ""
            with open(file_path, 'rb') as f:
                pdf_reader = PdfReader(f)
                for page in pdf_reader.pages:
                    text += page.extract_text()
            return text
        
        elif file_extension in ['.doc', '.docx']:
            doc = DocxDocument(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        else:
            print(f"Unsupported file type: {file_extension}")
            return ""
    
    except Exception as e:
        print(f"Error reading file {file_path}: {str(e)}")
        return ""

def read_files(file_paths):
    """
    Reads a series of files of various types, concatenates their contents,
    and returns a Document object with the combined content.
    
    Args:
        file_paths (list): A list of file paths to read.
    
    Returns:
        Document: A Document object with the combined contents of all the files.
    """
    combined_text = ""
    
    for file_path in file_paths:
        combined_text += read_file(file_path) + "\n"
    
    return Document(page_content=combined_text.strip())


